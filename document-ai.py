"""
RAG based - Document Intelligence - Company policies

Ingest many company documents (PDF, DOCX, TXT), build a FAISS vector index,
and run a retrieval-augmented Q&A system.

Run with mode="openai" or mode="hf_local".
- openai: uses OpenAIEmbeddings + OpenAI LLM (requires OPENAI_API_KEY)
- hf_local: uses HuggingFaceEmbeddings + a local HF model (no API key)
"""

import os
from pathlib import Path
from typing import List

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from tqdm import tqdm

# Document loaders
import PyPDF2
import docx

# Embeddings
from langchain.embeddings import HuggingFaceEmbeddings, OpenAIEmbeddings
from langchain.llms import HuggingFacePipeline, OpenAI

# transformers
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline

# ----------------------------
# Utility: load file content
# ----------------------------
def load_pdf(path: Path) -> str:
    text = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text.append(page_text)
    return "\n".join(text)

def load_docx(path: Path) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

# Supported extensions
LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_txt,
}


# Read documents from a folder
def load_documents_from_folder(folder: str) -> List[Document]:
    folder_path = Path(folder)
    assert folder_path.exists(), f"Folder not found: {folder}"
    docs: List[Document] = []

    file_paths = list(folder_path.rglob("*"))
    # keep only supported extensions
    file_paths = [p for p in file_paths if p.suffix.lower() in LOADERS]

    for p in tqdm(file_paths, desc="Loading files"):
        loader = LOADERS.get(p.suffix.lower())
        if not loader:
            continue
        try:
            text = loader(p)
            if not text or not text.strip():
                continue
            # Create a LangChain Document, preserve metadata (filename)
            docs.append(Document(page_content=text, metadata={"source": str(p)}))
        except Exception as e:
            print(f"Error loading {p}: {e}")
    return docs

# Chunk documents
def chunk_documents(docs: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    all_chunks: List[Document] = []
    for doc in docs:
        chunks = splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            metadata = dict(doc.metadata)
            metadata.update({"chunk": i})
            all_chunks.append(Document(page_content=chunk, metadata=metadata))
    return all_chunks

# Build FAISS index with embeddings
def build_vectorstore(chunks: List[Document], mode: str = "openai"):
    """
    mode: "openai" or "hf_local"
    returns: vectorstore
    """
    if mode == "openai":
        # Uses OpenAI embeddings (remote) - requires OPENAI_API_KEY env var
        emb = OpenAIEmbeddings()
    elif mode == "hf_local":
        # Uses a local sentence-transformers model (downloads once)
        emb = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    else:
        raise ValueError("mode must be 'openai' or 'hf_local'")

    texts = [d.page_content for d in chunks]
    metadatas = [d.metadata for d in chunks]
    vectorstore = FAISS.from_texts(texts, embedding=emb, metadatas=metadatas)
    return vectorstore

# RetrievalQA chain
def create_qa_chain(vectorstore, mode="openai"):
    retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 3})
    if mode == "openai":
        llm = OpenAI(temperature=0)  # requires OPENAI_API_KEY
    else:
        # Local HF LLM (seq2seq) - flan-t5-small is a good CPU-friendly choice
        model_name = "google/flan-t5-small"
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        hf_pipe = pipeline(
            "text2text-generation",
            model=model,
            tokenizer=tokenizer,
            max_length=256,
            truncation=True
        )
        llm = HuggingFacePipeline(pipeline=hf_pipe)
    qa_chain = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever, return_source_documents=True)
    return qa_chain

# Interactive QA loop
def interactive_loop(qa_chain):
    print("\nReady. Ask policy questions (type 'exit' to quit).\n")
    while True:
        q = input("Question: ").strip()
        if q.lower() in ("exit", "quit"):
            print("Bye.")
            break
        # Run the chain: returns answer and source docs
        result = qa_chain({"query": q})
        answer = result.get("result") or result.get("answer") or ""
        sources = result.get("source_documents") or []
        print("\n--- Answer ---\n")
        print(answer, "\n")
        if sources:
            print("--- Source chunks (top results) ---")
            for s in sources[:3]:
                src = s.metadata.get("source", "unknown")
                chunk_id = s.metadata.get("chunk", None)
                print(f"- Source: {src} (chunk: {chunk_id})")
            print()
        print("-----------\n")

# Main pipeline
def build_and_run(folder: str, mode: str = "openai"):
    # 1) Load raw documents
    raw_docs = load_documents_from_folder(folder)
    if not raw_docs:
        print("No documents found in folder (supported: .pdf, .docx, .txt). Exiting.")
        return

    # 2) Chunk documents
    chunks = chunk_documents(raw_docs, chunk_size=800, chunk_overlap=150)
    print(f"Total chunks: {len(chunks)}")

    # 3) Build vectorstore (this may take time for many docs)
    print("Building vectorstore (this may take a while for 100s of docs)...")
    vs = build_vectorstore(chunks, mode=mode)

    # 4) Create QA chain and run interactive loop
    qa_chain = create_qa_chain(vs, mode=mode)
    interactive_loop(qa_chain)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Policy Q&A (OpenAI or HF local)")
    parser.add_argument("--folder", type=str, required=True, help="Folder containing .pdf, .docx, .txt files")
    parser.add_argument("--mode", type=str, default="hf_local", choices=["openai", "hf_local"], help="Embedding/LLM mode")
    args = parser.parse_args()

    # If using openai mode, ensure OPENAI_API_KEY is set
    if args.mode == "openai" and "OPENAI_API_KEY" not in os.environ:
        raise RuntimeError("OPENAI_API_KEY is required for openai mode. Set environment variable and re-run.")

    build_and_run(args.folder, mode=args.mode)

